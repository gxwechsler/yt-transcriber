"""
Module: writers
Purpose: File writers for Markdown, Word, and JSON output
Created: 2026-01-28
Session: yt_trans_20260128_001 | Context: 1
"""
import json
from datetime import datetime
from pathlib import Path
from typing import Optional

from docx import Document as DocxDocument
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from src.core.models import VideoMeta, TranscriptEntry


def write_markdown(video: VideoMeta, output_path: Path) -> Path:
    """
    Write video transcript as Markdown file.
    
    Structure:
        - Title
        - Metadata block
        - Links (if present)
        - Chapters (if present)
        - Transcript with timestamps
    
    Returns:
        Path to written file
    """
    lines = [
        f"# {video.title}\n",
        f"**URL:** https://youtube.com/watch?v={video.video_id}  ",
        f"**Channel:** [{video.channel}]({video.channel_url})  ",
        f"**Subscribers:** {video.channel_follower_count:,}  ",
        f"**Date:** {video.upload_date_formatted}  ",
        f"**Duration:** {video.duration_formatted}  ",
        f"**Views:** {video.view_count:,} · **Likes:** {video.like_count:,}",
        ""
    ]
    
    # Links section
    if video.links:
        lines.append("\n## Links Mentioned\n")
        for link in video.links:
            ctx = f" — {link['context']}" if link.get('context') else ""
            lines.append(f"- <{link['url']}>{ctx}")
        lines.append("")
    
    # Chapters section
    if video.chapters:
        lines.append("\n## Chapters\n")
        for ch in video.chapters:
            t = int(ch.get('start_time', 0))
            timestamp = f"{t//60}:{t%60:02d}"
            lines.append(f"- **{timestamp}** — {ch.get('title', '')}")
        lines.append("")
    
    # Transcript section
    lines.append("\n---\n\n## Transcript\n")
    
    if video.transcript:
        current_ts = None
        para_lines = []
        
        for entry in video.transcript:
            ts = entry.timestamp if isinstance(entry, TranscriptEntry) else entry.get("timestamp", "")
            text = entry.text if isinstance(entry, TranscriptEntry) else entry.get("text", "")
            
            if ts != current_ts:
                if para_lines:
                    lines.append(f"**{current_ts}** {' '.join(para_lines)}\n")
                    para_lines = []
                current_ts = ts
            para_lines.append(text)
        
        # Flush remaining
        if para_lines:
            lines.append(f"**{current_ts}** {' '.join(para_lines)}\n")
    else:
        lines.append("*No transcript available*")
    
    output_path.parent.mkdir(parents=True, exist_ok=True)
    output_path.write_text('\n'.join(lines))
    
    return output_path


def write_docx(video: VideoMeta, output_path: Path) -> Path:
    """
    Write video transcript as Word document.
    
    Structure:
        - Title (centered heading)
        - Video Information section
        - Links section (if present)
        - Chapters section (if present)
        - Transcript (page break, timestamped paragraphs)
    
    Returns:
        Path to written file
    """
    doc = DocxDocument()
    
    # Title
    title_para = doc.add_heading(video.title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Metadata section
    doc.add_heading("Video Information", level=1)
    
    info_items = [
        f"URL: https://youtube.com/watch?v={video.video_id}",
        f"Channel: {video.channel}",
        f"Subscribers: {video.channel_follower_count:,}",
        f"Date: {video.upload_date_formatted}",
        f"Duration: {video.duration_formatted}",
        f"Views: {video.view_count:,}",
        f"Likes: {video.like_count:,}",
    ]
    
    for item in info_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(2)
    
    # Links section
    if video.links:
        doc.add_heading("Links Mentioned", level=1)
        for link in video.links:
            text = link['url']
            if link.get('context'):
                text += f" — {link['context']}"
            doc.add_paragraph(text, style='List Bullet')
    
    # Chapters section
    if video.chapters:
        doc.add_heading("Chapters", level=1)
        for ch in video.chapters:
            t = int(ch.get('start_time', 0))
            timestamp = f"{t//60}:{t%60:02d}"
            doc.add_paragraph(f"{timestamp} — {ch.get('title', '')}", style='List Bullet')
    
    # Transcript section
    doc.add_page_break()
    doc.add_heading("Transcript", level=1)
    
    if video.transcript:
        current_para_lines = []
        last_timestamp = None
        
        for entry in video.transcript:
            ts = entry.timestamp if isinstance(entry, TranscriptEntry) else entry.get("timestamp", "")
            text = entry.text if isinstance(entry, TranscriptEntry) else entry.get("text", "")
            
            if last_timestamp is None or ts != last_timestamp:
                # Flush current paragraph
                if current_para_lines:
                    p = doc.add_paragraph()
                    if last_timestamp:
                        ts_run = p.add_run(f"{last_timestamp} ")
                        ts_run.bold = True
                        ts_run.font.size = Pt(9)
                    p.add_run(' '.join(current_para_lines))
                    current_para_lines = []
                last_timestamp = ts
            
            current_para_lines.append(text)
        
        # Flush remaining
        if current_para_lines:
            p = doc.add_paragraph()
            if last_timestamp:
                ts_run = p.add_run(f"{last_timestamp} ")
                ts_run.bold = True
                ts_run.font.size = Pt(9)
            p.add_run(' '.join(current_para_lines))
    else:
        doc.add_paragraph("No transcript available", style='Intense Quote')
    
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    
    return output_path


def write_json(video: VideoMeta, output_path: Path) -> Path:
    """
    Write video metadata as JSON file.
    
    Contains structured data for programmatic use:
        - All metadata fields
        - Chapters
        - Links
        - Processing timestamp
    
    Returns:
        Path to written file
    """
    # Convert transcript entries to dicts if needed
    transcript_data = []
    for entry in video.transcript:
        if isinstance(entry, TranscriptEntry):
            transcript_data.append({"timestamp": entry.timestamp, "text": entry.text})
        else:
            transcript_data.append(entry)
    
    data = {
        "id": video.video_id,
        "title": video.title,
        "channel": video.channel,
        "channel_url": video.channel_url,
        "upload_date": video.upload_date_formatted,
        "duration": video.duration,
        "duration_formatted": video.duration_formatted,
        "view_count": video.view_count,
        "like_count": video.like_count,
        "description": video.description,
        "chapters": video.chapters,
        "links": video.links,
        "transcript_entries": len(transcript_data),
        "processed_at": datetime.now().isoformat(),
        # Naming metadata
        "saved_as": {
            "author": video.proposed_author,
            "topic": video.proposed_topic,
            "year": video.proposed_year
        }
    }
    
    output_path.parent.mkdir(parents=True, exist_ok=True)
    output_path.write_text(json.dumps(data, indent=2, ensure_ascii=False))
    
    return output_path
